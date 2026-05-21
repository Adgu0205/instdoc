import io
import docx
from app.parsers.txt import clean_text

def parse_docx(file_bytes: bytes) -> str:
    """
    Extracts text and tables from DOCX file bytes.
    """
    text_content = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        
        # Parse paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
                
        # Parse tables and merge as textual representations
        for table in doc.tables:
            for row in table.rows:
                # Deduplicate merged cells by creating a list of unique texts
                row_cells_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and (not row_cells_text or row_cells_text[-1] != cell_text):
                        row_cells_text.append(cell_text)
                if row_cells_text:
                    text_content.append(" | ".join(row_cells_text))
                    
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    full_text = "\n\n".join(text_content)
    cleaned = clean_text(full_text)
    
    if not cleaned:
        raise ValueError("DOCX file does not contain any readable text.")
        
    return cleaned
